"""网页与 Word 共享的 Markdown 渲染核心。"""

from __future__ import annotations

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .sop_layout import wrap_sop_semantic_layout

# nl2br：把正文中的单行换行视为硬换行（<br>），与编辑器「所见即所得」一致。
# 配方正文常逐行罗列「配料：用量」，缺少它时单换行会被 Markdown 当作空格合并成一行。
MARKDOWN_EXTENSIONS = ["tables", "md_in_html", "nl2br"]

ALLOWED_TAGS = {
    "h1", "h2", "h3", "h4", "h5", "h6",
    "p", "br", "hr",
    "strong", "b", "em", "i", "del", "s",
    "ul", "ol", "li",
    "table", "thead", "tbody", "tfoot", "tr", "th", "td",
    "blockquote", "pre", "code",
    "div", "a",
}

ALLOWED_ATTRS = {
    "div": {"class", "style"},
    "h3": {"class"},
    "a": {"href", "target", "rel"},
}

SAFE_PROTOCOLS = {"http:", "https:", "mailto:"}


def _sanitize_html(soup: BeautifulSoup) -> None:
    """递归剥离不在白名单中的标签与属性，防止 XSS。"""
    dirty_tags = list(soup.find_all(True))
    for tag in dirty_tags:
        if tag.name not in ALLOWED_TAGS:
            tag.unwrap()
            continue
        allowed = ALLOWED_ATTRS.get(tag.name, set())
        for attr in list(tag.attrs):
            if attr not in allowed:
                del tag[attr]
                continue
            if attr == "style":
                style_val = (tag[attr] or "").lower().replace(" ", "")
                if "page-break-after:always" in style_val:
                    tag[attr] = "page-break-after: always;"
                else:
                    del tag[attr]
            elif tag.name == "a" and attr == "href":
                href = (tag.get("href") or "").strip()
                if not href or href.startswith(("#", "/")):
                    tag[attr] = href
                elif not any(href.lower().startswith(proto) for proto in SAFE_PROTOCOLS):
                    del tag[attr]
            elif tag.name == "a" and attr == "target":
                if tag.get("target") != "_blank":
                    del tag[attr]
            elif tag.name == "a" and attr == "rel":
                tag[attr] = "noopener noreferrer"
        if tag.name == "a" and tag.get("target") == "_blank":
            tag["rel"] = "noopener noreferrer"


def _postprocess(html_fragment: str) -> str:
    wrapper = BeautifulSoup(f"<div>{html_fragment}</div>", "html.parser")
    root = wrapper.div
    if root is None:
        return html_fragment

    _sanitize_html(root)

    for div in root.find_all("div"):
        style = (div.get("style") or "").lower().replace(" ", "")
        if "page-break-after" in style and "always" in style:
            existing = div.get("class") or []
            if isinstance(existing, str):
                existing = [existing]
            div["class"] = list(existing) + ["sop-print-page-break"]

    for table in list(root.find_all("table")):
        parent = table.parent
        parent_classes = parent.get("class") if parent and parent.name else None
        if parent_classes and "table-scroll" in parent_classes:
            continue
        scroll = wrapper.new_tag("div", attrs={"class": "table-scroll"})
        table.replace_with(scroll)
        scroll.append(table)

    inner = root.decode_contents()
    return wrap_sop_semantic_layout(inner)


def render_markdown_to_html(markdown_text: str) -> str:
    html = markdown.markdown(markdown_text, extensions=MARKDOWN_EXTENSIONS)
    return _postprocess(html)


# ---------------------------------------------------------------------------
# Word (.docx) 渲染核心：与 md2docx 历史版式保持一致
# ---------------------------------------------------------------------------

FONT_NAME = "黑体"
FONT_NAME_EN = "Arial"
BODY_SIZE = 10
H1_SIZE = 16
H2_SIZE = 12
TABLE_HEADER_SIZE = 10
TABLE_BODY_SIZE = 10
QUOTE_SIZE = 10


def _set_run_font(run, size=BODY_SIZE, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = FONT_NAME_EN
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex,
    })
    shading.append(elem)


def _set_cell_margin(cell, top=20, bottom=20, left=40, right=40):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = tcPr.makeelement(qn("w:tcMar"), {})
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        e = tcMar.makeelement(qn(f"w:{side}"), {qn("w:w"): str(val), qn("w:type"): "dxa"})
        tcMar.append(e)
    tcPr.append(tcMar)


def _set_row_no_split(row, keep_with_next=False):
    """允许行内内容跨页，但表头与数据行尽量同页。"""
    if keep_with_next:
        for cell in row.cells:
            for para in cell.paragraphs:
                pPr = para._element.get_or_add_pPr()
                pPr.append(pPr.makeelement(qn("w:keepNext"), {}))


def _add_rich_text_to_cell(cell, html_content):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = Pt(11)
    soup = BeautifulSoup(html_content, "html.parser")

    def process_node(node, paragraph):
        if isinstance(node, NavigableString):
            text = str(node)
            if text.strip() or text == " ":
                run = paragraph.add_run(text)
                _set_run_font(run, size=TABLE_BODY_SIZE)
        elif node.name == "br":
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(11)
        elif node.name in ("strong", "b"):
            for child in node.children:
                if isinstance(child, NavigableString):
                    run = paragraph.add_run(str(child))
                    _set_run_font(run, size=TABLE_BODY_SIZE, bold=True)
                else:
                    paragraph = process_node(child, paragraph)
        elif node.name in ("em", "i"):
            for child in node.children:
                if isinstance(child, NavigableString):
                    run = paragraph.add_run(str(child))
                    _set_run_font(run, size=TABLE_BODY_SIZE)
                    run.italic = True
                else:
                    paragraph = process_node(child, paragraph)
        else:
            for child in node.children:
                paragraph = process_node(child, paragraph)
        return paragraph

    for child in soup.children:
        para = process_node(child, para)


def render_markdown_to_docx(markdown_text: str) -> Document:
    html = markdown.markdown(markdown_text, extensions=MARKDOWN_EXTENSIONS)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME_EN
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(BODY_SIZE)
    style.paragraph_format.space_before = Pt(1)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.line_spacing = Pt(11)

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    for element in soup.children:
        if isinstance(element, NavigableString):
            if element.strip():
                doc.add_paragraph(element.strip())
            continue

        if element.name == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(element.get_text())
            _set_run_font(run, size=H1_SIZE, bold=True)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
        elif element.name == "h2":
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            _set_run_font(run, size=H2_SIZE, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            p_border = p._element.get_or_add_pPr()
            bottom_border = p_border.makeelement(qn("w:pBdr"), {})
            bottom_elem = bottom_border.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "666666",
            })
            bottom_border.append(bottom_elem)
            p_border.append(bottom_border)
        elif element.name == "h3":
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            _set_run_font(run, size=9, bold=True)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
        elif element.name == "hr":
            continue
        elif element.name == "div":
            style_attr = element.get("style") or ""
            if "page-break" in style_attr:
                p = doc.add_paragraph()
                run = p.add_run()
                run._r.append(run._r.makeelement(qn("w:br"), {qn("w:type"): "page"}))
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            continue
        elif element.name == "table":
            rows = element.find_all("tr")
            if not rows:
                continue
            header_cells = rows[0].find_all(["th", "td"])
            num_cols = len(header_cells)
            if num_cols == 0:
                continue
            table = doc.add_table(rows=0, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
            tblW = tblPr.makeelement(qn("w:tblW"), {qn("w:w"): "5000", qn("w:type"): "pct"})
            tblPr.append(tblW)
            tblLayout = tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
            tblPr.append(tblLayout)
            col_width_pct = 5000 // num_cols
            hdr_row = table.add_row()
            for i, th in enumerate(header_cells):
                cell = hdr_row.cells[i]
                cell.text = ""
                tcPr = cell._element.get_or_add_tcPr()
                tcW = tcPr.makeelement(qn("w:tcW"), {qn("w:w"): str(col_width_pct), qn("w:type"): "pct"})
                tcPr.append(tcW)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(th.get_text().strip())
                _set_run_font(run, size=TABLE_HEADER_SIZE, bold=True)
                _set_cell_shading(cell, "E8E8E8")
                _set_cell_margin(cell, top=15, bottom=15, left=30, right=30)
            _set_row_no_split(hdr_row, keep_with_next=True)
            for tr in rows[1:]:
                tds = tr.find_all("td")
                row = table.add_row()
                for i, td in enumerate(tds):
                    if i < num_cols:
                        cell = row.cells[i]
                        tcPr = cell._element.get_or_add_tcPr()
                        tcW = tcPr.makeelement(qn("w:tcW"), {qn("w:w"): str(col_width_pct), qn("w:type"): "pct"})
                        tcPr.append(tcW)
                        inner_html = "".join(str(c) for c in td.children)
                        _add_rich_text_to_cell(cell, inner_html)
                        _set_cell_margin(cell, top=15, bottom=15, left=30, right=30)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(1)
            spacer.paragraph_format.space_after = Pt(1)
            spacer.paragraph_format.line_spacing = Pt(4)
            spacer.paragraph_format.keep_with_next = True
        elif element.name == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run(element.get_text().strip())
            _set_run_font(run, size=QUOTE_SIZE, bold=True, color=(180, 30, 30))
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
        elif element.name == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for child in element.children:
                if isinstance(child, NavigableString):
                    run = p.add_run(str(child))
                    _set_run_font(run, size=BODY_SIZE)
                elif child.name in ("strong", "b"):
                    run = p.add_run(child.get_text())
                    _set_run_font(run, size=BODY_SIZE, bold=True)
                elif child.name == "br":
                    p.add_run("\n")
                else:
                    run = p.add_run(child.get_text())
                    _set_run_font(run, size=BODY_SIZE)
        elif element.name in ("ul", "ol"):
            for li in element.find_all("li"):
                p = doc.add_paragraph(li.get_text(), style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

    return doc
